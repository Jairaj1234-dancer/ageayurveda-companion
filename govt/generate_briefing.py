"""Generate an internal "founder's eyes only" Word briefing on the entire
state of the Age Ayurveda Companion platform — what's built, the funding
research findings, the structural decisions, the 90-day plan, the open
questions. Includes family context where it's operationally useful (this
is internal; the no-family-names rule applies to external materials only).

Output: ~/Documents/AgeAyurveda/Internal-Briefing.docx
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor, Inches

OUT = Path.home() / "Documents" / "AgeAyurveda" / "Internal-Briefing.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

INK         = RGBColor(0x1B, 0x18, 0x14)
INK_SOFT    = RGBColor(0x4A, 0x41, 0x3A)
INK_FADED   = RGBColor(0x87, 0x7C, 0x70)
TERRACOTTA  = RGBColor(0xB8, 0x4A, 0x1F)
NEEM        = RGBColor(0x2E, 0x5D, 0x3F)
GOLD        = RGBColor(0xA8, 0x89, 0x5C)
INDIGO      = RGBColor(0x2C, 0x3E, 0x5C)
ASHOK       = RGBColor(0x1F, 0x4A, 0x8B)
HAIR        = RGBColor(0xD7, 0xCD, 0xB9)
PAPER       = RGBColor(0xF4, 0xEF, 0xE5)


def set_cell_background(cell, hex_color: str) -> None:
    """Set table cell background colour."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, top=True, bottom=True, left=False, right=False, color="D7CDB9", sz="8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge, on in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single" if on else "nil")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_horizontal_rule(doc, color="D7CDB9", height_pt=1):
    """Insert a slim horizontal rule between paragraphs."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(height_pt * 8)))  # 1/8 pt units
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_styles(doc: Document) -> None:
    """Define a coherent type system that matches the project's print
    aesthetic — Cambria for body (close to Fraunces in feel and bundled
    with Word on every desktop), Calibri for sans tag/eyebrow lines, and
    Consolas for monospace. Heading sizes scale proportionally."""
    styles = doc.styles

    # Body
    normal = styles["Normal"]
    normal.font.name = "Cambria"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    # Headings
    for level, size, color in [
        ("Heading 1", 24, INK),
        ("Heading 2", 17, INK),
        ("Heading 3", 13, INK_SOFT),
    ]:
        h = styles[level]
        h.font.name = "Cambria"
        h.font.size = Pt(size)
        h.font.bold = False
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(20 if level == "Heading 1" else 14)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True

    # Custom styles
    if "Eyebrow" not in [s.name for s in styles]:
        s = styles.add_style("Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Calibri"
        s.font.size = Pt(8)
        s.font.color.rgb = TERRACOTTA
        s.font.all_caps = True
        s.font.bold = True
        s.paragraph_format.space_after = Pt(2)
        s.paragraph_format.space_before = Pt(0)

    if "Lead" not in [s.name for s in styles]:
        s = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Cambria"
        s.font.size = Pt(13)
        s.font.italic = True
        s.font.color.rgb = INK_SOFT
        s.paragraph_format.space_after = Pt(10)
        s.paragraph_format.line_spacing = 1.4

    if "Mono" not in [s.name for s in styles]:
        s = styles.add_style("Mono", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Consolas"
        s.font.size = Pt(10)
        s.font.color.rgb = INK_SOFT
        s.paragraph_format.space_after = Pt(4)

    if "Sidebar" not in [s.name for s in styles]:
        s = styles.add_style("Sidebar", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Cambria"
        s.font.size = Pt(10)
        s.font.color.rgb = INK_SOFT
        s.font.italic = True
        s.paragraph_format.left_indent = Cm(0.6)
        s.paragraph_format.space_after = Pt(4)


def add_para(doc, text, style="Normal", bold=False, color=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color is not None:
        run.font.color.rgb = color
    return p


def add_eyebrow(doc, text, color=TERRACOTTA):
    p = doc.add_paragraph(style="Eyebrow")
    run = p.add_run(text)
    run.font.color.rgb = color


def add_kv_line(doc, key, value):
    """Bold key, normal value — single paragraph 'Key — value' line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    k = p.add_run(f"{key} — ")
    k.bold = True
    k.font.color.rgb = INK
    v = p.add_run(value)
    v.font.color.rgb = INK_SOFT


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    return p


def add_pullquote(doc, text, attr=None):
    """Hairline-bordered pull quote, indented."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # Top + bottom hairline
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "8")
        b.set(qn("w:color"), "A8895C")
        p_bdr.append(b)
    p_pr.append(p_bdr)
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.italic = True
    run.font.color.rgb = INK
    if attr:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.8)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(10)
        r = p2.add_run("— " + attr)
        r.font.size = Pt(9)
        r.font.color.rgb = INK_FADED


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------

doc = Document()
setup_styles(doc)

# Page setup — A4-ish margins, slightly tighter for density
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)

# Footer with confidentiality + page number
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp_run = fp.add_run("Internal · for the founder's eyes only · ")
fp_run.font.size = Pt(8)
fp_run.font.color.rgb = INK_FADED
fp_run.italic = True

# Page number field
page_run = fp.add_run()
page_run.font.size = Pt(8)
page_run.font.color.rgb = INK_FADED
fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
page_run._r.append(fldChar1); page_run._r.append(instrText); page_run._r.append(fldChar2)

# ===========================================================================
# COVER
# ===========================================================================

cover_p = doc.add_paragraph(style="Eyebrow")
cover_p.add_run("Confidential · for the founder's eyes only").font.color.rgb = TERRACOTTA

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(24)
title.paragraph_format.space_after = Pt(6)
t_run = title.add_run("Age Ayurveda Companion")
t_run.font.name = "Cambria"
t_run.font.size = Pt(34)
t_run.font.color.rgb = INK
sub = title.add_run("\nState of the platform · plan of action · funding pathway")
sub.font.name = "Cambria"
sub.font.size = Pt(15)
sub.italic = True
sub.font.color.rgb = INK_SOFT

add_horizontal_rule(doc, color="A8895C", height_pt=2)

# Cover meta block
meta = doc.add_paragraph()
meta.paragraph_format.space_before = Pt(8)
meta_run = meta.add_run(
    f"Compiled · {datetime.now().strftime('%B %Y')}\n"
    "Author · Jairaj Sharma (founder) with research synthesis from three parallel AI agents\n"
    "Distribution · solo. Do not share.\n"
    "Sister documents · pitch.html (investor) · govt/pitch.html (government) · "
    "govt/action-plan.md · govt/ccras-letter.md"
)
meta_run.font.name = "Consolas"
meta_run.font.size = Pt(9)
meta_run.font.color.rgb = INK_FADED

doc.add_page_break()

# ===========================================================================
# TABLE OF CONTENTS
# ===========================================================================

doc.add_heading("Contents", level=1)
toc = [
    ("1.", "The picture in one minute"),
    ("2.", "What is actually running today"),
    ("3.", "The strategic thesis"),
    ("4.", "Government funding research — full findings"),
    ("5.", "The structural decision (the load-bearing call)"),
    ("6.", "The 90-day plan of action"),
    ("7.", "Capital efficiency forecast"),
    ("8.", "Risk register"),
    ("9.", "Where this fits in the broader Nitya / Baidyanath stack"),
    ("10.", "Open questions and decisions you owe yourself"),
    ("11.", "Reference — schemes, links, contacts"),
]
for num, title_ in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    n = p.add_run(num + "  ")
    n.font.color.rgb = TERRACOTTA
    n.bold = True
    t = p.add_run(title_)
    t.font.color.rgb = INK
doc.add_page_break()

# ===========================================================================
# 1 · ONE MINUTE
# ===========================================================================

doc.add_heading("1 · The picture in one minute", level=1)
add_eyebrow(doc, "Read this if nothing else")

add_para(doc,
    "You have built — and have running in production — a citation-disciplined "
    "AI chat platform that grounds every answer in classical Sanskrit Ayurvedic "
    "verses. Backend, frontend, multi-tenant authentication, billing telemetry, "
    "rate limits, 55 passing tests. Two classical sources live, hybrid retrieval, "
    "tool-calling for product recommendations, streaming with mid-stream product "
    "cards. This is no longer a research project. It is a deployable platform.",
    style="Lead",
)

add_para(doc,
    "The strategic thesis is a collision of three forces — a 2,500-year-old "
    "classical corpus, a 2026 AI capable of grounded reasoning over it, and a "
    "100-year heritage that vouches for the institutional credibility of the work. "
    "No other party in India has all three at once."
)

add_para(doc,
    "On the funding side, three parallel research agents have surfaced ~₹70–80 L "
    "of conservative non-dilutive Year-1 money plus 70%+ compute discounts — but "
    "all of it gates on one structural decision: incorporate the AI platform as a "
    "sister Pvt Ltd, NOT as a Nitya subsidiary. DPIIT explicitly disqualifies "
    "subsidiaries. Get this wrong and Sec 80-IAC, FFS, and the IndiaAI startup "
    "track are foreclosed permanently."
)

add_pullquote(doc,
    "The single highest-asymmetry move is a research partnership letter to CCRAS. "
    "Every other AI startup has to cold-pitch the council. You can walk the letter in.",
    attr="The conclusion of the entire funding research"
)

add_para(doc, "What you owe yourself this month:", bold=True)
add_bullet(doc, "Engage your CA on the structural decision (sister Pvt Ltd vs subsidiary). Pre-incorporation paperwork must show no 'splitting up' of an existing business.")
add_bullet(doc, "Send the CCRAS research-partnership letter (drafted, ready). Heritage letterhead, co-signed by Dhananjay Sharma's office for institutional weight.")
add_bullet(doc, "Lock an academic co-PI for the AYUSH EMR grant (MNNIT Allahabad CS dept is the natural pick).")
add_bullet(doc, "Decide whether the AI work being done now inside Nitya is material enough to trigger 'reconstruction' risk. If yes, get a CA-signed defensibility memo before incorporating the new entity.")

doc.add_page_break()

# ===========================================================================
# 2 · WHAT IS RUNNING
# ===========================================================================

doc.add_heading("2 · What is actually running today", level=1)
add_eyebrow(doc, "The platform · operational state")

add_para(doc,
    "Repository · ~/Projects/ageayurveda-companion. Backend: FastAPI + SQLAlchemy "
    "async, SQLite for dev, Postgres-ready. Frontend: TypeScript + Vite IIFE "
    "widget, Shadow-DOM scoped, embeds in any Shopify store. LLM layer: "
    "Anthropic Claude Opus 4.7 with adaptive thinking and prompt caching wired."
)

doc.add_heading("2.1 · Backend grounding pipeline", level=2)
add_bullet(doc, "Corpus — 34 verses across Aṣṭāṅga Hṛdaya Sūtrasthāna (20) and Charaka Saṃhitā Sūtrasthāna (14). Each verse stored with Sanskrit mūla, Roman transliteration, English paraphrase, structured metadata.")
add_bullet(doc, "Retrieval — hybrid BM25 + dense semantic, fused via Reciprocal Rank Fusion. Multilingual MiniLM embeddings. Stop-word-filtered, Devanagari-aware tokenizer (this caught two real bugs — Devanagari combining marks and citation regex no-comma case).")
add_bullet(doc, "Generation — frozen system prompt with strict citation format. Every claim resolves to a verse. Citation regex recognises Charaka, Aṣṭāṅga Hṛdaya, Suśruta, Bhāvaprakāśa formats.")
add_bullet(doc, "Tool — recommend_age_ayurveda_products fires on wellness intent. Concern × dosha matcher over the live product catalog. Returns ranked products with reason text.")
add_bullet(doc, "Streaming — SSE with mid-stream tool_use events so the widget renders product cards before the final text completes streaming. Abortable mid-flight.")

doc.add_heading("2.2 · Multi-tenant platform", level=2)
add_bullet(doc, "Tenants table with public api_key (ageak_…), BYO Anthropic key, source allowlist, per-tenant rate limit, soft-revoke flag.")
add_bullet(doc, "Bearer token auth on grounded endpoints. Tenant resolves the Anthropic key (BYO billing — we don't carry inference cost), the source filter, and the rate limit dynamically.")
add_bullet(doc, "Admin CRUD endpoints — create, list, get, patch, rotate-key, revoke, usage telemetry, conversation listing.")
add_bullet(doc, "Per-tenant cost telemetry — daily token aggregates with cost estimate using the live model pricing table.")

doc.add_heading("2.3 · Frontend widget", level=2)
add_bullet(doc, "Toggle-able grounded mode (rule-based fallback when ANTHROPIC_API_KEY isn't configured — no behaviour change for existing deployments).")
add_bullet(doc, "Free-text input + canned-menu navigation. Streaming bubble with blinking cursor, mid-stream product cards, citation chips at end of message.")
add_bullet(doc, "Stop-streaming button (AbortController) and conversation reset (clears local conversation_id + transcript). 41 KB / 14 KB gzipped.")

doc.add_heading("2.4 · Test suite", level=2)
add_para(doc,
    "55 tests passing in 0.6 seconds. Coverage: BM25 tokenization + stop words + scoring + RRF "
    "(11), citation extractor (7), product tool dispatch + ranking (9), retrieval source filter (7), "
    "grounded chat tool loop with mocked client (5), tenant model + auth resolution + rate limiter (16). "
    "The suite caught two production bugs during construction — Devanagari combining marks "
    "splitting words and citations without commas mis-parsing the source name."
)

# Snapshot table
doc.add_heading("2.5 · Snapshot of current state", level=2)
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.autofit = True
hdr = table.rows[0].cells
hdr[0].text = "Surface"
hdr[1].text = "State"
for k in hdr:
    for run in k.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = INK

snapshot = [
    ("Backend grounded chat", "Tenant-aware, BYO billing, source-filterable, hybrid retrieval, tool-using, streaming, abortable"),
    ("Corpus", "34 verses · Aṣṭāṅga Hṛdaya + Charaka Saṃhitā Sūtrasthāna · schema-locked YAML"),
    ("Frontend widget", "Streaming · stop + reset + free-text input · citation chips · 14 KB gzipped"),
    ("Multi-tenant", "Tenants table + auth + per-tenant rate + per-tenant cost + admin CRUD"),
    ("Tests", "55 passing in 0.6s"),
    ("Pitches", "pitch.html (investor) · govt/pitch.html (government) · all family names stripped per your preference"),
    ("Plans", "govt/action-plan.md · govt/ccras-letter.md (ready to send)"),
]
for row_data in snapshot:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    for c in row:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_page_break()

# ===========================================================================
# 3 · STRATEGIC THESIS
# ===========================================================================

doc.add_heading("3 · The strategic thesis", level=1)
add_eyebrow(doc, "Why this platform · why us · why now", color=GOLD)

add_para(doc, "Three forces, one moment.", style="Lead")

add_para(doc,
    "Each force on its own is interesting. The collision is the company. Nobody "
    "else owns all three at once — and you already do."
)

doc.add_heading("3.1 · The texts (2,500 years)", level=2)
add_para(doc,
    "Charaka Saṃhitā, Suśruta Saṃhitā, Aṣṭāṅga Hṛdaya, Bhāvaprakāśa — the original "
    "Sanskrit mūla pāṭha is firmly in the public domain. No party has packaged it "
    "as a structured, verse-tagged retrieval surface that AI systems can use as "
    "ground-truth. The CCRAS-Berhampur and CCRAS-Anuvadini MoUs of 2025 confirm "
    "the policy direction is right; the execution gap is wide open."
)

doc.add_heading("3.2 · The model (2026)", level=2)
add_para(doc,
    "Twelve months ago this was a research project. Today it ships against "
    "production-grade APIs with adaptive thinking, prompt caching, and a 1M-token "
    "context window. Multilingual classical-text reasoning with citation "
    "discipline is now operationally tractable. The window between 'good enough' "
    "and 'commoditised' is open today; it will not be open in three years."
)

doc.add_heading("3.3 · The house (1917)", level=2)
add_para(doc,
    "Baidyanath. Three generations. Trust no Series-A founder can buy. The "
    "credibility — and, equally, the existing relationships across the AYUSH "
    "ministry, CCRAS, the All India Ayurveda Manufacturers Association — is "
    "what every other AI wellness startup will be missing for at least a decade. "
    "Dhananjay Sharma's standing as President of Baidyanath is a strategic asset "
    "for govt-engagement that the platform can leverage without diluting either "
    "the heritage entity or the new research arm."
)

doc.add_heading("3.4 · The three revenue surfaces", level=2)
add_para(doc, "One engine, three markets. Same grounding layer, three commercial doors:")
add_bullet(doc, "B2B API license — every Ayurveda D2C brand (Kapiva, Forest Essentials, Wow Skin, Mamaearth Ayurveda lines, etc.). ₹50K–2L/month per brand. Tenant model already shipping; ~3,000 SMB buyers globally.")
add_bullet(doc, "Consumer subscription app — pocket Ayurvedic guide, Prakṛti tracking, dincharya, telemedicine handoff. ₹199–499/month. 350M+ wellness consumer market.")
add_bullet(doc, "Clinical white-label — for AYUSH-affiliated hospitals, integrative-medicine clinics, Patanjali clinics, Kerala Ayurveda chains. ₹2L–10L annually per institution.")

add_pullquote(doc,
    "Build the infrastructure layer for the world's oldest medicine. We have the texts. "
    "We have the heritage. The model is finally good enough. The market is fragmented and "
    "ready. This is the moment when one company defines the category."
)

doc.add_page_break()

# ===========================================================================
# 4 · GOVT FUNDING RESEARCH FINDINGS
# ===========================================================================

doc.add_heading("4 · Government funding research — full findings", level=1)
add_eyebrow(doc, "From three parallel research agents · April 2026")

add_para(doc,
    "Three agents ran in parallel: one on the IndiaAI Mission compute scheme, "
    "one on MSME (Udyam) + DPIIT Startup India recognition specific to your "
    "structural situation, one on adjacent AI/digital-health schemes (MeitY, "
    "AYUSH, NITI, NASSCOM, state policies). Findings consolidated below."
)

doc.add_heading("4.1 · IndiaAI Mission Compute Capacity", level=2)
add_kv_line(doc, "Status (Apr 2026)", "Portal live; 38,000+ GPUs onboarded vs. original 10,000 target. 14 service providers empanelled.")
add_kv_line(doc, "Two tracks", "(a) 40% subsidy for general DPIIT-recognised startups + MSMEs; (b) 100% subsidy reserved for foundational-model builders (Sarvam, Gnani, Gan.ai, Soket so far).")
add_kv_line(doc, "Pricing post-subsidy", "H100 ~₹92/hr (Vaishnaw's quoted floor: ₹67/hr). 70–80% cheaper than AWS/Azure list.")
add_kv_line(doc, "Eligibility gates", "DPIIT recognition required for startup track; ≥51% Indian shareholding; re-verification every 2 years.")
add_kv_line(doc, "Best providers for our workload", "E2E Networks, Yotta, or Jio — H100/A100 inventory at lowest published rates with managed Kubernetes.")
add_kv_line(doc, "Application portal", "indiaai.gov.in/hub/indiaai-compute-capacity; compute booking via PM Gati Shakti subdomain (still on staging — minor red flag).")
add_kv_line(doc, "Login methods", "DigiLocker, e-Pramaan, or Jan Parichay")

add_para(doc, "The honest verdict for our specific situation:", bold=True, color=TERRACOTTA)
add_para(doc,
    "The 40% IndiaAI subsidy does not help our current architecture — Anthropic charges "
    "per token, not per GPU-hour. The lever is real only for self-hosted workloads (embeddings, "
    "fine-tuning a small classical-text model, batch indexing). Switching the live chat to "
    "self-hosted Llama-3-70B is a 4–8 week engineering lift with measurable quality regression, "
    "and self-hosting only beats Claude API economics above ~10–20M tokens/day sustained — "
    "below that, Claude API is cheaper once you factor staff time and the quality gap."
)
add_para(doc,
    "Where IndiaAI does make sense: corpus expansion. We will need to embed ~10,000 verses "
    "as the corpus grows. Doing that on subsidised H100s at ₹92/hr is genuinely cheap. "
    "Apply for the 40% track post-DPIIT, scope the workload as 'classical-text embedding "
    "and batch indexing,' and reserve the option to fine-tune a small Sanskrit-aware model later."
)
add_para(doc,
    "Skip the 100% foundational-model track. It requires open-sourcing the model and "
    "repositioning as a 'sovereign Ayurveda LLM' project — not the right trade for "
    "Companion's current architecture, but worth holding in reserve.",
    color=INK_SOFT,
)

# ----- 4.2 MSME / DPIIT
doc.add_heading("4.2 · MSME (Udyam) + DPIIT Startup India", level=2)

add_para(doc, "MSME thresholds (revised April 2025):")
add_bullet(doc, "Micro — ≤₹2.5 Cr investment / ≤₹10 Cr turnover")
add_bullet(doc, "Small — ≤₹25 Cr / ≤₹100 Cr")
add_bullet(doc, "Medium — ≤₹125 Cr / ≤₹500 Cr")
add_para(doc, "Software is registered as a Service enterprise; same caps apply across services and manufacturing post-2020 unification.")

add_para(doc, "NIC codes for the AI platform:")
add_bullet(doc, "62011 — Custom software development (closest fit for Companion)")
add_bullet(doc, "62013 — Software support and maintenance")
add_bullet(doc, "62091 — Other IT services")
add_para(doc, "All three on a single Udyam registration. Multiple codes per Udyam are explicitly allowed via the 'Add More' field.")

add_para(doc, "DPIIT Startup India key facts:")
add_bullet(doc, "Eligibility — Pvt Ltd / LLP / registered partnership; <10 yrs old; turnover ≤₹100 Cr; not formed by splitting or reconstruction of existing business; working on innovation/scalability.")
add_bullet(doc, "Sec 80-IAC tax holiday — 100% profit deduction for any 3 of first 10 years. Window extended by Budget 2025 to entities incorporated before 1 April 2030.")
add_bullet(doc, "Angel tax (Sec 56(2)(viib)) — abolished from AY 2025-26 by Finance Act 2024 for ALL unlisted companies. The DPIIT-specific carve-out is now historical.")
add_bullet(doc, "FFS — ₹10,000 Cr corpus deployed via SIDBI to SEBI-registered AIFs which then invest in DPIIT startups. Indirect — relevant only at Series A.")

add_para(doc, "MSME benefits worth pursuing:")
add_bullet(doc, "CGTMSE collateral-free loan ceiling raised to ₹10 Cr (from ₹5 Cr) in 2025 — useful working-capital backstop.")
add_bullet(doc, "GeM Startup Runway — EMD/turnover/experience waivers; software/SaaS categories exist; realistic only if there's a public-sector buyer (AYUSH ministry, state telemedicine).")
add_bullet(doc, "45-day buyer-payment rule (Sec 43B(h)) — useful only for B2B sales to profit-making Indian buyers; marginal for D2C/Shopify-embedded chatbot use case.")

add_para(doc, "UP State perks (Prayagraj is in Purvanchal, the most-incentivised zone):")
add_bullet(doc, "Stamp duty — 100% exemption on premises purchase")
add_bullet(doc, "Capital subsidy — 15-25% (Purvanchal); +2% bump for women/SC-ST")
add_bullet(doc, "Interest subvention — 50% of interest for 5 years, cap ₹25 L/unit")
add_bullet(doc, "UP IT & ITeS Policy 2022 — additional EPF reimbursement, lease-rental subsidy, patent reimbursement")

# ----- 4.3 Adjacent
doc.add_heading("4.3 · Adjacent AI / digital-health schemes", level=2)

add_para(doc, "AYUSH Ministry — strongest content-fit:", bold=True, color=NEEM)
add_bullet(doc, "AYUSH CCRAS — already signed MoUs with Berhampur University (palm-leaf manuscript digitisation) and Anuvadini AI (13-language Ayurveda translation). Live MoU appetite. The ask is non-monetary; what we offer is open-licensed corpus + benchmark + toolkit + papers.")
add_bullet(doc, "AYUSH ANUDAN portal — Extra-Mural Research (EMR) grants up to ~₹30 L per project, non-dilutive. Private orgs eligible but grants skew academic; lock an academic co-PI (MNNIT Allahabad, IIIT-A) before applying.")
add_bullet(doc, "NAMASTE portal — standardised ASU codes/terminologies for Electronic Health Records. Useful as integration target, not money source.")
add_bullet(doc, "National AYUSH Mission (NAM) — flows to states via UP State AYUSH Society. Mostly clinical/infrastructure, weak fit unless productising for govt clinics.")

add_para(doc, "MeitY non-IndiaAI:", bold=True)
add_bullet(doc, "TIDE 2.0 — 51+ TBI-routed grants for deeptech/AI. EIR ₹4L stipend + grant up to ~₹25L. IIIT-D, IIM-A Ventures, NIT-RKL, IIIT-D as TBI partners. Worth applying as fallback if SAMRIDH doesn't land.")
add_bullet(doc, "SAMRIDH — accelerator with 1:1 matching investment up to ₹40L. 2025 Cohort II is digital-health focused at C-CAMP Bangalore. Best fit on the MeitY side.")
add_bullet(doc, "IndiaAI Governance Sandbox (Nov 2025 guidelines) — regulatory pathway for healthcare AI, not a grant. Track only.")

add_para(doc, "NITI Aayog:", bold=True)
add_bullet(doc, "AIRAWAT (now AIRAWAT-PSAI) — 82-node DGX-A100 cluster operational at C-DAC Pune. Distinct from IndiaAI Mission's GPU subsidy. ~65 startups using it. Access via NASSCOM Open Source GenAI Grand Challenge (HPAIC track). Faster route to actual GPU than waiting on IndiaAI tender.")

add_para(doc, "State-level — UP vs alternatives:", bold=True)
add_bullet(doc, "UP StartinUP — DPIIT-recognised startups get ₹17.5K/month sustenance × 12 months + ₹7.5L seed × 1.5 (Purvanchal multiplier) = ₹13L total. AI is a notified priority sector with a CoE at IIIT-Allahabad.")
add_bullet(doc, "Karnataka ELEVATE NxT — up to ₹1 Cr deeptech grant; ₹518 Cr Startup Policy 2025-30 with explicit AI focus. Compliance cost of a Bangalore branch is ~₹1-2L/yr; worth it for the ₹1 Cr ceiling if other avenues exhaust.")
add_bullet(doc, "Telangana T-Fund + T-AIM — ₹25L-₹1Cr; most active AI-specific state body.")
add_bullet(doc, "Verdict — don't move HQ. Baidyanath + UP heritage IS the moat. A single Bangalore branch office for ELEVATE NxT eligibility is the only state-arbitrage that pays off.")

add_para(doc, "Health tech / NHA:", bold=True)
add_bullet(doc, "ABDM Sandbox — 919 health-tech innovators integrated; free API + certification; not a grant but unlocks ABHA/HPR/HFR integration. Strategic surface for clinical SKU.")
add_bullet(doc, "SAHI + BODH (launched Mar 2026) — national health-AI strategy + IIT-Kanpur/NHA open benchmarking platform. Companion's citation-traceable architecture is designed to be benchmarked. Track weekly; onboarding mechanics still publishing.")

add_para(doc, "Skip:", bold=True, color=INK_FADED)
add_bullet(doc, "DLI / SPECS — semiconductor design only")
add_bullet(doc, "MeghRaj / NIC Cloud — government tenants only")
add_bullet(doc, "PARAM Utkarsh / NSM — academic-host gating; AIRAWAT-PSAI is the better-fit AI compute path")

doc.add_page_break()

# ===========================================================================
# 5 · STRUCTURAL DECISION
# ===========================================================================

doc.add_heading("5 · The structural decision (the load-bearing call)", level=1)
add_eyebrow(doc, "The single decision that gates everything else", color=TERRACOTTA)

add_para(doc, "There are three options. Only one is right.", style="Lead")

# Decision matrix
table = doc.add_table(rows=4, cols=5)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Option"
hdr[1].text = "DPIIT eligible?"
hdr[2].text = "Sec 80-IAC?"
hdr[3].text = "FFS?"
hdr[4].text = "Verdict"
for c in hdr:
    for run in c.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = INK
    set_cell_background(c, "ECE4D2")

rows = [
    ("(a) Extend Nitya Udyam · add software NIC codes", "No (>10 yrs)", "No", "No", "Cheapest. Forfeits ~₹50L+. Skip."),
    ("(b) Wholly-owned subsidiary of Nitya", "NO — DPIIT bars subsidiaries", "No", "No", "AVOID."),
    ("(c) Sister Pvt Ltd · founder/family direct shareholders", "Yes", "Yes", "Yes", "THE PLAY."),
]
for i, row_data in enumerate(rows):
    row = table.rows[i + 1].cells
    for j, val in enumerate(row_data):
        row[j].text = val
        for p in row[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
        if j == 4:
            for run in row[j].paragraphs[0].runs:
                run.font.bold = True
                if "AVOID" in val:
                    run.font.color.rgb = TERRACOTTA
                elif "THE PLAY" in val:
                    run.font.color.rgb = NEEM
                else:
                    run.font.color.rgb = INK_FADED

doc.add_heading("5.1 · The rule that drives this", level=2)
add_para(doc,
    "DPIIT G.S.R. 127(E) Sec 2(d) explicitly disqualifies (a) holding/subsidiary "
    "structures and (b) entities formed by 'splitting up or reconstruction of "
    "an existing business'. A wholly-owned subsidiary of Nitya is therefore not "
    "DPIIT-eligible at all, and DPIIT additionally checks combined group turnover."
)
add_para(doc,
    "The 'splitting up' bar has been litigated narrowly under Sec 33B IT Act. "
    "Pre-revenue or exploratory work is generally defensible. Booked AI revenue "
    "or capitalised AI assets inside Nitya before incorporation creates real risk."
)

doc.add_heading("5.2 · What the sister Pvt Ltd structure unlocks", level=2)
add_bullet(doc, "Fresh Udyam under software NIC codes (62011, 62013, 62091) — Micro / services classification, well within ₹10 Cr turnover ceiling for years.")
add_bullet(doc, "DPIIT recognition (free, ~10 working days at startupindia.gov.in).")
add_bullet(doc, "Sec 80-IAC eligibility — 100% profit deduction, any 3 of first 10 years.")
add_bullet(doc, "FFS access at Series A.")
add_bullet(doc, "IndiaAI startup-track compute subsidy.")
add_bullet(doc, "UP StartinUP Purvanchal seed + sustenance grants.")

doc.add_heading("5.3 · Compliance pre-work (before incorporation)", level=2)
add_bullet(doc, "CA opinion — review of any AI-related expenditure or capitalised assets currently inside Nitya. Document this is exploratory / non-revenue. If material AI revenue has been booked, get a CA-signed memo on 'reconstruction' defensibility.")
add_bullet(doc, "Draft minimal IP-licence + services agreement template for arm's-length transactions between Nitya and the new entity. No asset spin-out. Distinct customers.")
add_bullet(doc, "Reserve company name via MCA RUN (e.g. 'Age Ayurveda Companion Pvt Ltd' or similar — confirm availability).")
add_bullet(doc, "Decide shareholding pattern. The cleanest is founder + 1-2 family individuals as direct shareholders (NOT Nitya as parent). Discuss with family — Sanjay, Dhananjay, Paranjay — whether they want stakes directly. If they do, structure their holding as individuals rather than via Nitya to preserve the structure.")

add_pullquote(doc,
    "Get this wrong and Sec 80-IAC, FFS, the IndiaAI startup track, and the UP "
    "StartinUP grants are foreclosed permanently. Get it right and the entire "
    "central + state startup stack opens up.",
    attr="The structural decision in one sentence"
)

doc.add_page_break()

# ===========================================================================
# 6 · 90-DAY PLAN
# ===========================================================================

doc.add_heading("6 · The 90-day plan of action", level=1)
add_eyebrow(doc, "Concrete · sequenced · ready to execute")

# Gantt as a monospace block
add_para(doc, "Critical path:", bold=True)
add_para(doc,
    "Incorporation → DPIIT → 80-IAC + IndiaAI Compute. Everything downstream "
    "blocks on DPIIT. Parallel-able: CCRAS letter — start drafting Week 1, send "
    "Week 4 (no DPIIT dependency). Slowest decision: 80-IAC at ~120 days."
)

p = doc.add_paragraph(style="Mono")
p.paragraph_format.space_before = Pt(8)
gantt = (
    "Week           1   2   3   4   5   6   7   8   9  10  11  12\n"
    "─────────────────────────────────────────────────────────\n"
    "CA opinion     ████\n"
    "Incorporation       ████\n"
    "Udyam                       ██\n"
    "DPIIT                       ████\n"
    "80-IAC application                 ████ → decision ~W18\n"
    "GST + bank                       ██\n"
    "CCRAS letter draft                  ██\n"
    "CCRAS letter sent                       █\n"
    "EMR co-PI outreach                     ████\n"
    "EMR proposal draft                            ████\n"
    "EMR submission                                    ██\n"
    "SAMRIDH (next cohort)                              ████ →\n"
    "UP StartinUP                                  ██\n"
    "IndiaAI Compute                                   ████\n"
    "NASSCOM DeepTech                       ██\n"
    "AIRAWAT-PSAI                                    ████\n"
)
p.add_run(gantt).font.size = Pt(8.5)

doc.add_heading("6.1 · Weeks 1–2 · Foundations", level=2)
add_bullet(doc, "Engage CA. Review AI work currently inside Nitya. Get defensibility opinion if any revenue has been booked.")
add_bullet(doc, "Family conversation — Sanjay Sharma, Dhananjay Sharma, Paranjay Sharma — on whether the new entity is founder-only or includes other family individuals as direct shareholders. Avoid Nitya-as-parent.")
add_bullet(doc, "Reserve company name (MCA RUN form).")
add_bullet(doc, "Incorporate sister Pvt Ltd via SPICe+. Two directors min. PAN + TAN auto-generated.")
add_bullet(doc, "Draft IP-licence + services agreement template for the Nitya ↔ new entity relationship.")

doc.add_heading("6.2 · Weeks 3–4 · Registrations", level=2)
add_bullet(doc, "File Udyam under NIC 62011 / 62013 / 62091. Free. Same day.")
add_bullet(doc, "Apply DPIIT Startup India recognition. Free. ~10 working days.")
add_bullet(doc, "Apply Sec 80-IAC concurrently — the decision takes 120 days, start it early.")
add_bullet(doc, "GST registration (gst.gov.in). 7 working days.")
add_bullet(doc, "Bank current account opening — ICICI Privilege or HDFC Smartup. Razorpay account for SaaS billing.")
add_bullet(doc, "Draft and dispatch CCRAS research-partnership letter — heritage letterhead, co-signed by Dhananjay's office. Send by Speed Post AND email. (Letter is ready: govt/ccras-letter.md.)")
add_bullet(doc, "Open accounts on anudan.ayush.gov.in, startinup.up.gov.in, startupindia.gov.in.")

doc.add_heading("6.3 · Weeks 5–8 · Apply", level=2)
add_bullet(doc, "Identify and approach academic co-PI for AYUSH EMR grant — MNNIT Allahabad CS dept dean is the natural opening. Frame as research partnership on classical-text retrieval.")
add_bullet(doc, "Draft full EMR proposal: 'AI-Augmented Retrieval and Provenance Validation of Classical Ayurvedic Texts'. Lead PI = academic; co-PI = founder; industry partner = Baidyanath.")
add_bullet(doc, "MeitY SAMRIDH application to C-CAMP digital-health cohort (or wait for next cohort opening).")
add_bullet(doc, "NASSCOM DeepTech Club membership + Open-Source GenAI Grand Challenge cohort application → unlocks AIRAWAT-PSAI compute.")

doc.add_heading("6.4 · Weeks 9–12 · Stack the rest", level=2)
add_bullet(doc, "UP StartinUP Purvanchal seed grant application post-DPIIT.")
add_bullet(doc, "IndiaAI Compute application via indiaai.gov.in/hub/indiaai-compute-capacity. Scope: classical-text embedding + batch indexing workloads.")
add_bullet(doc, "CGTMSE-backed working-capital line at preferred bank — useful backstop, not primary funding.")
add_bullet(doc, "Track SAHI/BODH onboarding mechanics for validation pathway publication.")
add_bullet(doc, "Submit EMR proposal via ANUDAN portal (Week 10-11).")

doc.add_page_break()

# ===========================================================================
# 7 · CAPITAL EFFICIENCY FORECAST
# ===========================================================================

doc.add_heading("7 · Capital efficiency forecast", level=1)
add_eyebrow(doc, "Conservative · Year-1 stack")

table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Source"
hdr[1].text = "Form"
hdr[2].text = "Amount"
hdr[3].text = "Probability"
for c in hdr:
    for run in c.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)
    set_cell_background(c, "ECE4D2")

forecast = [
    ("CCRAS Research MoU", "Non-monetary, credentialing", "—", "High"),
    ("AYUSH EMR Grant", "Non-dilutive grant", "~₹25 L", "Medium (with academic co-PI)"),
    ("MeitY SAMRIDH", "1:1 matching grant", "~₹30 L", "Medium"),
    ("UP StartinUP — Purvanchal", "Seed + sustenance", "~₹13 L", "High (post-DPIIT auto)"),
    ("IndiaAI Compute subsidy", "In-kind GPU credit", "~₹20–40 L equivalent", "High"),
    ("80-IAC tax holiday", "Future tax shield × 3 yrs", "Deferred", "High"),
    ("CGTMSE collateral-free loan", "Working-capital backstop", "Up to ₹10 Cr ceiling", "High (post-Udyam)"),
    ("Conservative Y1 cash + in-kind", "", "~₹70–80 L", ""),
    ("Aggressive (all + UP MSME stamp/capital subsidies + Karnataka ELEVATE NxT)", "", "~₹1.2–1.5 Cr", ""),
]
for row_data in forecast:
    row = table.add_row().cells
    for j, val in enumerate(row_data):
        row[j].text = val
        for p in row[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                if "Conservative" in row_data[0] or "Aggressive" in row_data[0]:
                    run.font.bold = True
                    if "Conservative" in row_data[0]:
                        run.font.color.rgb = NEEM
                    else:
                        run.font.color.rgb = GOLD

add_para(doc,
    "These numbers exclude the moat compounding from a CCRAS partnership credential — "
    "which is the unmeasured value driver. They also exclude the platform's own commercial "
    "revenue (B2B tenant licensing, ~₹1.5 L/mo per brand at the mid-tier price point).",
    style="Sidebar"
)

doc.add_page_break()

# ===========================================================================
# 8 · RISK REGISTER
# ===========================================================================

doc.add_heading("8 · Risk register", level=1)
add_eyebrow(doc, "What can go wrong · how we mitigate", color=TERRACOTTA)

table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Risk"
hdr[1].text = "Mitigation"
for c in hdr:
    for run in c.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)
    set_cell_background(c, "ECE4D2")

risks = [
    ("'Splitting up' doctrine bites at DPIIT review",
     "CA paper-trail of pre-incorporation AI work as exploratory, non-revenue. No asset spin-out from Nitya. Document distinct customers from day one."),
    ("EMR grant rejected for lacking academic credentials",
     "Lock co-PI early (MNNIT Allahabad / IIIT-A). Frame as research instrument, not product. Pre-coordinate with CCRAS scholarly contacts."),
    ("IndiaAI End-User Policy disallows foreign LLM API use in product",
     "Workload separation — IndiaAI compute used only for embeddings + fine-tuning; Anthropic API stays as the chat-time inference layer. The policy PDF is currently unreadable; read it before applying."),
    ("SAMRIDH cohort timing misses",
     "TIDE 2.0 fallback (₹25 L) via IIIT-D / IIM-A Ventures partner TBIs."),
    ("80-IAC profit deduction unusable in early years (no profits yet)",
     "Defer claim to Y3-Y5 — rule allows any 3 consecutive years out of first 10. Just file the application early; usage is at our discretion."),
    ("UP StartinUP processing delays",
     "Apply concurrently with DPIIT; expect 60-120 day lag. Not on critical path."),
    ("CCRAS letter goes unanswered",
     "Follow-up via existing AYUSH ministry contacts after 3 weeks. Use Dhananjay Sharma's institutional standing as Baidyanath President to warm-call CCRAS leadership directly. Do NOT skip the formal letter — it establishes precedence."),
    ("Family disagreement on entity structure or shareholding",
     "Pre-discussion with Sanjay, Dhananjay, Paranjay before incorporation. Position the sister Pvt Ltd as a research arm of the family group, not as a competitive entity. Their direct shareholding (vs. Nitya holding) preserves DPIIT eligibility — frame this as a regulatory necessity, not a governance preference."),
    ("Quality regression if forced to migrate off Claude API",
     "Don't migrate. Use IndiaAI compute for embedding/indexing only. Keep Claude API as the inference layer. The economics support this — Claude is cheaper than self-hosted below ~10-20M tokens/day."),
    ("Pre-existing bcrypt + Python 3.13 incompatibility (caught during testing)",
     "Pinned bcrypt < 4.1 in requirements.txt. Documented in code. Confirmed working."),
    ("Schema-drift on Postgres deployment",
     "lifespan auto-create handles new tables for SQLite dev only. Run alembic revision --autogenerate before any prod deploy. Currently the corpus_chunks + tenants + retrievals column on messages need migrations."),
]
for risk, mit in risks:
    row = table.add_row().cells
    row[0].text = risk
    row[1].text = mit
    for c in row:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_page_break()

# ===========================================================================
# 9 · WHERE THIS FITS
# ===========================================================================

doc.add_heading("9 · Where this fits in the broader Nitya / Baidyanath stack", level=1)
add_eyebrow(doc, "Strategic placement", color=GOLD)

add_para(doc,
    "Nitya Naturals is the contract-manufacturing arm of the Baidyanath group — "
    "GMP-certified, EU-export-focused, ~₹1 Cr current revenue with a ₹5 Cr target. "
    "Plant utilisation is at 20-30%; lead generation is essentially Baidyanath "
    "referrals + word of mouth, no active outbound. The seven internal software "
    "projects you've been building are mostly operational tooling. Companion is "
    "the only one with platform shape — the only one with a path to outgrow its parent."
)

doc.add_heading("9.1 · The seven Nitya software projects", level=2)
add_bullet(doc, "Client Portal — internal, B2B catalog + ordering")
add_bullet(doc, "Production Planning — replacement for spreadsheet scheduling")
add_bullet(doc, "Lead Tracking CRM — internal lead management")
add_bullet(doc, "Age Ayurveda Companion — THIS — the platform play, the moonshot, the only one with $1B shape")
add_bullet(doc, "Website Revamp — brand")
add_bullet(doc, "Shipping Cost Calculator — operational utility")
add_bullet(doc, "Reorder Prediction Dashboard — internal demand forecasting")

doc.add_heading("9.2 · Revenue ladder", level=2)
add_kv_line(doc, "Today", "Nitya at ~₹1 Cr · Companion pre-revenue")
add_kv_line(doc, "2027 target", "Nitya at ₹5 Cr · Companion onboarding first 10 paying tenants at ~₹1.5L/month each = ~₹15L MRR")
add_kv_line(doc, "2028-29", "Nitya at ₹15-20 Cr (with active EU lead-gen) · Companion at ₹50 Cr ARR (250 tenants × ₹1.5L/month + consumer subscription line)")
add_kv_line(doc, "Long-arc", "Companion as $1B category-leader exit. Nitya as the cash-generating brand + manufacturing backbone that funded it.")

doc.add_heading("9.3 · Why the structure matters operationally", level=2)
add_para(doc,
    "Keeping the AI platform legally separate from Nitya isn't just a tax-optimisation. "
    "It's a strategic separation. The platform serves multiple Ayurveda brands as B2B "
    "tenants — including, eventually, brands that compete with AGE Ayurveda. The "
    "operational neutrality of the platform is what makes it sellable to those brands. "
    "If Companion is legally Nitya, Patanjali is never going to license it. If "
    "Companion is a sister entity that sells infrastructure to everyone, Patanjali "
    "becomes a plausible tenant."
)
add_para(doc,
    "The same logic applies to investors at Series A. They're investing in an "
    "AI infrastructure company, not in a manufacturing subsidiary. The cleaner the "
    "separation, the higher the valuation multiple."
)

doc.add_page_break()

# ===========================================================================
# 10 · OPEN QUESTIONS
# ===========================================================================

doc.add_heading("10 · Open questions and decisions you owe yourself", level=1)
add_eyebrow(doc, "What still needs your call", color=TERRACOTTA)

decisions = [
    ("Shareholding pattern of the new Pvt Ltd",
     "Founder-only? Founder + spouse? Founder + Sanjay/Dhananjay/Paranjay as individual shareholders (NOT Nitya)? Each option has different DPIIT, family-governance, and exit-economics implications. Decide before incorporation."),
    ("Family alignment on the spin-out",
     "Have Sanjay, Dhananjay, and Paranjay been briefed that the AI work is moving to a separate entity? This needs a sit-down conversation, not an email. Frame as a regulatory necessity and as a strategic move for the family — Companion is the moonshot, Nitya stays the cash engine, both win."),
    ("CCRAS warm intro path",
     "Does Dhananjay's office have an existing relationship with the current CCRAS Director General? If yes, route the letter via that channel for institutional weight. If no, send formal letter cold and follow up via AYUSH ministry contacts at the All India Ayurvedic Manufacturers Association level."),
    ("Co-PI for the EMR grant",
     "MNNIT Allahabad CS dept is the closest geographical pick. IIIT-Allahabad is also viable. Banaras Hindu U Faculty of Ayurveda is a wildcard with deeper Ayurvedic credentials. Decide who to approach first — and via whom."),
    ("Self-host migration question",
     "Are you ever going to migrate off Claude API to self-hosted? At current volumes the answer is no — Claude is cheaper. At scale (>10M tokens/day) the answer becomes yes. Set the threshold; don't migrate prematurely."),
    ("Open-source posture",
     "The proposed CCRAS deliverables include an open-licensed corpus + benchmark + toolkit. This is the right move for credibility. But it does mean someone could fork the toolkit and compete. The moat isn't the code — it's the corpus + heritage. Be comfortable with this before signing the MoU."),
    ("Bangalore branch for Karnataka ELEVATE NxT",
     "Worth ~₹1-2 L/yr in compliance for a ₹1 Cr grant ceiling. Decide based on whether the AYUSH/UP stack lands first. Don't pursue prematurely."),
    ("Consumer app GTM timing",
     "The B2B-API surface is current focus. The consumer subscription app is Phase 5 (2027). Resist the temptation to build it earlier — it's a different product, different team, different acquisition motion. Stay disciplined."),
    ("Production deployment",
     "When does Companion actually go live in front of paying tenants? Backend is ready. Widget is ready. Schema migrations need to be written for Postgres. ANTHROPIC_API_KEY needs to be set. Probably 1-2 weeks of runway work + a soft launch with one friendly tenant. Book it."),
    ("Memory hygiene on the spin-out",
     "If the new Pvt Ltd uses any Nitya assets (servers, code, IP), document the licence terms. The arm's-length agreement is template-able; have the CA paper-trail it before any value moves between entities."),
]
for q, body in decisions:
    add_para(doc, q, bold=True, color=INK)
    add_para(doc, body, color=INK_SOFT)

doc.add_page_break()

# ===========================================================================
# 11 · REFERENCE
# ===========================================================================

doc.add_heading("11 · Reference — schemes, links, contacts", level=1)
add_eyebrow(doc, "Bookmarks for the road")

doc.add_heading("11.1 · Application portals", level=2)
add_kv_line(doc, "Udyam (MSME)", "udyamregistration.gov.in")
add_kv_line(doc, "DPIIT Startup India", "startupindia.gov.in")
add_kv_line(doc, "Sec 80-IAC", "startupindia.gov.in/content/sih/en/form80iac.html")
add_kv_line(doc, "GST", "gst.gov.in")
add_kv_line(doc, "MCA SPICe+", "mca.gov.in")
add_kv_line(doc, "AYUSH ANUDAN (EMR)", "anudan.ayush.gov.in")
add_kv_line(doc, "IndiaAI Compute", "indiaai.gov.in/hub/indiaai-compute-capacity")
add_kv_line(doc, "IndiaAI Cloud Computing Portal", "staging2.pmgatishakti.gov.in/IndiaAICompute")
add_kv_line(doc, "AIRAWAT-PSAI", "airawat.cdac.in")
add_kv_line(doc, "MeitY SAMRIDH (C-CAMP)", "ccamp.res.in/digital-health-platform-samridh")
add_kv_line(doc, "MeitY TIDE 2.0", "msh.meity.gov.in/schemes/tide")
add_kv_line(doc, "UP StartinUP funding", "startinup.up.gov.in/funding")
add_kv_line(doc, "NASSCOM DeepTech Club", "nasscom.in/deeptech")
add_kv_line(doc, "ABDM Sandbox", "abdm.gov.in")

doc.add_heading("11.2 · Key contacts to develop", level=2)
add_kv_line(doc, "CCRAS Director General", "Janakpuri, New Delhi · ccrasdg-ayush@gov.in (verify before sending)")
add_kv_line(doc, "MNNIT Allahabad CS Dept", "Dean office · approach for EMR co-PI")
add_kv_line(doc, "IIIT-Allahabad", "AI CoE under UP IT-ITeS Policy 2022 · also EMR co-PI option")
add_kv_line(doc, "C-CAMP Bangalore", "SAMRIDH digital-health cohort applications")
add_kv_line(doc, "FITT-IIT Delhi", "Parallel SAMRIDH cohort")

doc.add_heading("11.3 · Repository state", level=2)
add_kv_line(doc, "Codebase", "~/Projects/ageayurveda-companion")
add_kv_line(doc, "Investor pitch", "pitch.html")
add_kv_line(doc, "Government pitch", "govt/pitch.html")
add_kv_line(doc, "Action plan (this is the markdown source)", "govt/action-plan.md")
add_kv_line(doc, "CCRAS letter (drafted, ready to send)", "govt/ccras-letter.md")
add_kv_line(doc, "Backend tests", "55 passing in 0.6s · backend/tests/")

doc.add_heading("11.4 · The single sentence", level=2)
add_pullquote(doc,
    "If only one thing happens this month, the CCRAS letter goes out. Every other "
    "lever benefits from having a CCRAS engagement in motion, and no other AI startup "
    "in India is in a position to walk that letter in.",
    attr=None
)

# Final note
add_horizontal_rule(doc, color="A8895C", height_pt=2)
add_para(doc, "End of briefing.", style="Eyebrow")
final = doc.add_paragraph()
final_run = final.add_run(
    "This document was assembled from three parallel research agents, the live "
    "codebase, the investor pitch deck, the government pitch deck, the action plan, "
    "and the CCRAS letter. It is intentionally comprehensive — read it once cover-to-"
    "cover, then keep it as a reference. The action items in §1 (\"What you owe yourself "
    "this month\") are the operational handle. Everything else is depth and rationale."
)
final_run.font.size = Pt(10)
final_run.font.color.rgb = INK_FADED
final_run.italic = True

# Save
doc.save(str(OUT))
print(f"Briefing written: {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")
